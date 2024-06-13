import streamlit as st
from docx import Document
from dotenv import load_dotenv
from PyPDF2 import PdfReader
from streamlit_extras.add_vertical_space import add_vertical_space
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from openai import OpenAI
import os
import re

# Load the .env file
load_dotenv()

api_key = os.getenv('OPENAI_API_KEY')

client = OpenAI(api_key=api_key)

def analyze_text_with_llm(text, client):
    prompt = (
        "You are an expert in the Abuse Prevention team at Intuit, "
        "reviewing a product requirement document (PRD) written by MailChimp PMs. "
        "Read the following text and identify any potential abuse cases. "
        "For each abuse case, suggest appropriate control methods. "
        "Text: {text}"
    ).format(text=text)
    
    response = client.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are an expert in the Abuse Prevention team at Intuit."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=1500,
        n=1,
        stop=None,
        temperature=0.7,
    )
    
    return response.choices[0].message.content.strip()

# Extracts text from a PDF file
def extract_text_from_pdf(pdf):
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Extracts text from a Docx file
def extract_text_from_docx(docx):
    doc = Document(docx)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    return text

def clean_text(text):
    # Remove non-XML compatible characters
    text = re.sub(r'[^\x09\x0A\x0D\x20-\x7F]+', '', text)
    return text

def main():
    st.header("PRD Analyzer 🚨")
    add_vertical_space(2)
    
    # File uploader for PDF and Word documents
    uploaded_file = st.file_uploader("Upload File", type=['pdf', 'docx'])
 
    if uploaded_file is not None:
        try:
            # Extract text based on file type
            if uploaded_file.type == "application/pdf":
                text = extract_text_from_pdf(uploaded_file)
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = extract_text_from_docx(uploaded_file)
            else:
                st.error("Unsupported file type. Please upload a PDF or Word document.")
                return
                
            if not text:
                st.error("Unable to extract text from the uploaded document. Please check the file and try again.")
                return
            
            text = clean_text(text)
            store_name = uploaded_file.name.rsplit('.', 1)[0]
            st.write(f'Processing file: {store_name}')
    
            # Load or create FAISS vector store
            if os.path.exists(f"{store_name}.faiss"):
                VectorStore = FAISS.load_local(store_name)
                st.write('Embeddings Loaded from the Disk')
            else:
                embeddings = OpenAIEmbeddings(openai_api_key=api_key)
                text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200,
                    length_function=len
                )
                chunks = text_splitter.split_text(text=text)
                VectorStore = FAISS.from_texts(chunks, embedding=embeddings)
                VectorStore.save_local(store_name)
    
            findings = analyze_text_with_llm(text, client)
            
            st.subheader("Potential Abuse Cases and Control Methods:")
            st.write(findings)

            # Display options for each recommendation
            recommendations = findings.split('\n\n') 
            new_doc.add_heading('PRD Document', 0)
            new_doc.add_paragraph(text)

            new_doc.add_heading('Appendix: Abuse Cases and Recommendations', 1)
            for rec in recommendations:
                st.write(rec)
                col1, col2, col3 = st.columns(3)
                with col1:
                    if st.button('Create a Jira Tix', key=f'jira_{rec[:30]}'):
                        st.write('Jira ticket creation simulated.')
                with col2:
                    if st.button('Auto Fix', key=f'autofix_{rec[:30]}'):
                        new_doc.add_paragraph(rec)
                with col3:
                    if st.button('Skip', key=f'skip_{rec[:30]}'):
                        st.write('Skipped recommendation.')
            
            st.subheader("Export Modified PRD Document")
            if st.button('Export Document'):
                new_doc_path = f"{store_name}_modified.docx"
                new_doc.save(new_doc_path)
                with open(new_doc_path, 'rb') as f:
                    st.download_button('Download Modified PRD', f, file_name=new_doc_path)
                
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

if __name__ == '__main__':
    main()